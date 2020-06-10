from docx import Document

document = Document()
document.add_heading('Test', 0)

p = document.add_paragraph("Let's see if this works")

table = document.add_table(rows=1, cols=2)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Number'
hdr_cells[1].text = 'Even'
for i in range(1000):
    row_cells = table.add_row().cells
    row_cells[0].text = str(i)
    row_cells[1].text = str(2*i)

document.save('test.docx')